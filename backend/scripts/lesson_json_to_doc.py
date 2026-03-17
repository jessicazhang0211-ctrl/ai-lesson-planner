import argparse
import json
import os
from datetime import datetime
from pathlib import Path
from typing import Any


def format_key(key: str) -> str:
    text = str(key).replace("_", " ").strip()
    if not text:
        return ""
    return text[0].upper() + text[1:]


def to_markdown_lines(value: Any, title: str = "Lesson Plan JSON") -> list[str]:
    lines: list[str] = [f"# {title}", ""]

    def walk(node: Any, level: int, name: str | None = None) -> None:
        if name:
            heading = "#" * min(level, 6)
            lines.append(f"{heading} {format_key(name)}")
            lines.append("")

        if isinstance(node, dict):
            if not node:
                lines.append("(empty object)")
                lines.append("")
                return
            for k, v in node.items():
                if isinstance(v, (dict, list)):
                    walk(v, min(level + 1, 6), k)
                else:
                    lines.append(f"- **{format_key(k)}**: {v}")
            lines.append("")
            return

        if isinstance(node, list):
            if not node:
                lines.append("(empty list)")
                lines.append("")
                return

            all_scalar = all(not isinstance(x, (dict, list)) for x in node)
            if all_scalar:
                for item in node:
                    lines.append(f"- {item}")
                lines.append("")
                return

            for idx, item in enumerate(node, 1):
                item_name = f"Item {idx}"
                if isinstance(item, dict):
                    maybe_phase = item.get("phase")
                    if isinstance(maybe_phase, str) and maybe_phase.strip():
                        item_name = maybe_phase.strip()
                walk(item, min(level + 1, 6), item_name)
            return

        lines.append(str(node))
        lines.append("")

    walk(value, 2)
    return lines


def write_markdown(json_data: dict, out_md: Path) -> None:
    title = json_data.get("lesson_title") or json_data.get("topic") or "Lesson Plan"
    md_lines = to_markdown_lines(json_data, title=str(title))
    out_md.write_text("\n".join(md_lines), encoding="utf-8")


def write_docx(json_data: dict, out_docx: Path) -> bool:
    try:
        from docx import Document
    except Exception:
        return False

    doc = Document()
    title = json_data.get("lesson_title") or json_data.get("topic") or "Lesson Plan"
    doc.add_heading(str(title), level=0)

    def walk(node: Any, level: int, name: str | None = None) -> None:
        if name:
            doc.add_heading(format_key(name), level=min(level, 4))

        if isinstance(node, dict):
            for k, v in node.items():
                if isinstance(v, (dict, list)):
                    walk(v, min(level + 1, 4), k)
                else:
                    p = doc.add_paragraph()
                    p.add_run(f"{format_key(k)}: ").bold = True
                    p.add_run(str(v))
            return

        if isinstance(node, list):
            if not node:
                doc.add_paragraph("(empty list)")
                return

            all_scalar = all(not isinstance(x, (dict, list)) for x in node)
            if all_scalar:
                for item in node:
                    doc.add_paragraph(str(item), style="List Bullet")
                return

            for idx, item in enumerate(node, 1):
                item_name = f"Item {idx}"
                if isinstance(item, dict):
                    maybe_phase = item.get("phase")
                    if isinstance(maybe_phase, str) and maybe_phase.strip():
                        item_name = maybe_phase.strip()
                walk(item, min(level + 1, 4), item_name)
            return

        doc.add_paragraph(str(node))

    walk(json_data, 1)
    doc.save(str(out_docx))
    return True


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert lesson-plan JSON to markdown/docx.")
    parser.add_argument("input", help="Path to lesson JSON file")
    parser.add_argument("--outdir", default="backend/exports", help="Output directory")
    parser.add_argument("--basename", default="", help="Output file base name (without extension)")
    parser.add_argument("--md-only", action="store_true", help="Only export markdown")
    args = parser.parse_args()

    in_path = Path(args.input)
    if not in_path.exists():
        raise FileNotFoundError(f"Input file not found: {in_path}")

    with in_path.open("r", encoding="utf-8") as f:
        payload = json.load(f)

    if not isinstance(payload, dict):
        raise ValueError("Input JSON must be an object")

    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    base = args.basename.strip() or in_path.stem
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_md = outdir / f"{base}_{ts}.md"
    out_docx = outdir / f"{base}_{ts}.docx"

    write_markdown(payload, out_md)
    print(f"[OK] Markdown exported: {out_md}")

    if args.md_only:
        return

    ok = write_docx(payload, out_docx)
    if ok:
        print(f"[OK] Word exported: {out_docx}")
    else:
        print("[WARN] python-docx not installed, skipped .docx export")
        print("       Install with: pip install python-docx")


if __name__ == "__main__":
    main()
